"""Document-related tools for reading and generating documents."""

import os
import re
import uuid
import time
from pathlib import Path

from langchain_core.tools import tool
from app.core.logging import get_logger

logger = get_logger(__name__)

# Temporary document storage — cleaned up periodically
DOCS_DIR = Path("/tmp/agentverse_docs")
DOCS_DIR.mkdir(parents=True, exist_ok=True)

# Track generated documents for cleanup
_generated_docs: dict[str, dict] = {}  # file_id -> {path, created_at, filename}


@tool
async def analyze_document(document_text: str) -> str:
    """Analyze a document and provide a structured analysis including key points,
    themes, and a summary. Use this when asked to analyze, review, or break down
    a document.

    Args:
        document_text: The full text content of the document to analyze.
    """
    logger.info("tool.analyze_document", text_length=len(document_text))

    if not document_text.strip():
        return "Error: No document text provided."

    # Truncate very long documents
    if len(document_text) > 50000:
        document_text = document_text[:50000] + "\n\n[Document truncated at 50,000 characters]"

    word_count = len(document_text.split())
    char_count = len(document_text)
    line_count = document_text.count("\n") + 1

    return (
        f"Document Statistics:\n"
        f"- Words: {word_count:,}\n"
        f"- Characters: {char_count:,}\n"
        f"- Lines: {line_count:,}\n\n"
        f"Document Content:\n{document_text}\n\n"
        f"Please provide your analysis of this document."
    )


@tool
async def extract_key_points(document_text: str, max_points: int = 10) -> str:
    """Extract key points, facts, figures, and dates from a document.
    Returns them as a structured list.

    Args:
        document_text: The text content to extract key points from.
        max_points: Maximum number of key points to extract (default: 10).
    """
    logger.info("tool.extract_key_points", text_length=len(document_text))

    if not document_text.strip():
        return "Error: No text provided for key point extraction."

    return (
        f"Please extract up to {max_points} key points from the following document:\n\n"
        f"{document_text[:30000]}\n\n"
        f"Format each key point clearly with a bullet point."
    )


# ── Format converters ─────────────────────────────────────

def _write_markdown(content: str, filepath: Path, title: str) -> None:
    """Write plain markdown content."""
    filepath.write_text(content, encoding="utf-8")


def _write_text(content: str, filepath: Path, title: str) -> None:
    """Write plain text content."""
    filepath.write_text(content, encoding="utf-8")


def _write_html(content: str, filepath: Path, title: str) -> None:
    """Write HTML-wrapped content."""
    html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; max-width: 800px; margin: 2rem auto; padding: 0 1rem; color: #1a1a2e; line-height: 1.7; }}
        h1, h2, h3 {{ color: #16213e; }}
        pre {{ background: #f0f0f5; padding: 1rem; border-radius: 8px; overflow-x: auto; }}
        code {{ background: #f0f0f5; padding: 2px 6px; border-radius: 4px; }}
        table {{ border-collapse: collapse; width: 100%; }}
        th, td {{ border: 1px solid #ddd; padding: 8px 12px; text-align: left; }}
        th {{ background: #f5f5fa; }}
    </style>
</head>
<body>
{content}
</body>
</html>"""
    filepath.write_text(html_content, encoding="utf-8")


def _write_docx(content: str, filepath: Path, title: str) -> None:
    """Convert markdown-style content to a Word .docx document."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Add title
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Parse markdown-like content into paragraphs
    lines = content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Headings
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        # Bullet points
        elif stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:]
            # Bold handling within bullet
            para = doc.add_paragraph(style="List Bullet")
            _add_formatted_text(para, text)
        # Numbered items
        elif re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s", "", stripped)
            para = doc.add_paragraph(style="List Number")
            _add_formatted_text(para, text)
        # Horizontal rule
        elif stripped in ("---", "***", "___"):
            doc.add_paragraph("─" * 50)
        # Regular paragraph
        else:
            para = doc.add_paragraph()
            _add_formatted_text(para, stripped)

        i += 1

    doc.save(str(filepath))


def _add_formatted_text(paragraph, text: str) -> None:
    """Parse simple markdown bold/italic in text and add formatted runs."""
    # Split on bold markers (**text**)
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def _write_xlsx(content: str, filepath: Path, title: str) -> None:
    """Convert content to an Excel spreadsheet.

    Tries to detect table-like data (CSV rows, markdown tables, tab-separated).
    Falls back to putting each line in a row if no structure is detected.
    """
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

    wb = Workbook()
    ws = wb.active
    ws.title = title[:31]  # Excel sheet name max 31 chars

    lines = [l for l in content.split("\n") if l.strip()]

    # Try to detect markdown table format (| col1 | col2 |)
    table_lines = [l for l in lines if "|" in l and not l.strip().startswith("|--") and not re.match(r"^\|[-:| ]+\|$", l.strip())]
    separator_lines = [l for l in lines if re.match(r"^\|[-:| ]+\|$", l.strip())]

    if len(table_lines) >= 2 and len(separator_lines) >= 1:
        # Markdown table detected
        rows = []
        for line in table_lines:
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            rows.append(cells)
        _fill_worksheet(ws, rows)
    else:
        # Try CSV-style (comma-separated with consistent column counts)
        csv_rows = [l.split(",") for l in lines if "," in l]
        if len(csv_rows) >= 2 and len(set(len(r) for r in csv_rows)) <= 2:
            rows = [[c.strip() for c in r] for r in csv_rows]
            _fill_worksheet(ws, rows)
        else:
            # Try tab-separated
            tsv_rows = [l.split("\t") for l in lines if "\t" in l]
            if len(tsv_rows) >= 2:
                rows = [[c.strip() for c in r] for r in tsv_rows]
                _fill_worksheet(ws, rows)
            else:
                # Fallback: each line is a row with one cell
                for idx, line in enumerate(lines, 1):
                    stripped = line.strip().lstrip("#").lstrip("- ").lstrip("* ").strip()
                    ws.cell(row=idx, column=1, value=stripped)
                    # Bold for lines that look like headers
                    if line.strip().startswith("#") or idx == 1:
                        ws.cell(row=idx, column=1).font = Font(bold=True, size=12)

    # Auto-adjust column widths
    for col in ws.columns:
        max_length = 0
        col_letter = col[0].column_letter
        for cell in col:
            try:
                if cell.value:
                    max_length = max(max_length, len(str(cell.value)))
            except Exception:
                pass
        ws.column_dimensions[col_letter].width = min(max_length + 4, 60)

    wb.save(str(filepath))


def _fill_worksheet(ws, rows: list[list[str]]) -> None:
    """Fill a worksheet with rows, styling the first row as a header."""
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )

    for row_idx, row in enumerate(rows, 1):
        for col_idx, value in enumerate(row, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.border = thin_border
            cell.alignment = Alignment(wrap_text=True, vertical="top")
            if row_idx == 1:
                cell.fill = header_fill
                cell.font = header_font


def _write_pdf(content: str, filepath: Path, title: str) -> None:
    """Convert content to a PDF document."""
    from fpdf import FPDF

    class PDF(FPDF):
        def header(self):
            self.set_font("Helvetica", "B", 10)
            self.set_text_color(120, 120, 120)
            self.cell(0, 8, title, align="R")
            self.ln(10)

        def footer(self):
            self.set_y(-15)
            self.set_font("Helvetica", "I", 8)
            self.set_text_color(150, 150, 150)
            self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", align="C")

    pdf = PDF()
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    # Title
    pdf.set_font("Helvetica", "B", 20)
    pdf.set_text_color(30, 30, 60)
    pdf.cell(0, 15, title, new_x="LMARGIN", new_y="NEXT")
    pdf.ln(5)

    # Horizontal rule after title
    pdf.set_draw_color(200, 200, 200)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(8)

    lines = content.split("\n")
    for line in lines:
        stripped = line.strip()

        if not stripped:
            pdf.ln(4)
            continue

        # Headings
        if stripped.startswith("### "):
            pdf.set_font("Helvetica", "B", 12)
            pdf.set_text_color(50, 50, 80)
            pdf.ln(3)
            pdf.multi_cell(0, 7, stripped[4:])
            pdf.ln(2)
        elif stripped.startswith("## "):
            pdf.set_font("Helvetica", "B", 14)
            pdf.set_text_color(40, 40, 70)
            pdf.ln(4)
            pdf.multi_cell(0, 8, stripped[3:])
            pdf.ln(3)
        elif stripped.startswith("# "):
            pdf.set_font("Helvetica", "B", 16)
            pdf.set_text_color(30, 30, 60)
            pdf.ln(5)
            pdf.multi_cell(0, 9, stripped[2:])
            pdf.ln(4)
        # Bullet points
        elif stripped.startswith("- ") or stripped.startswith("* "):
            pdf.set_font("Helvetica", "", 11)
            pdf.set_text_color(40, 40, 40)
            text = stripped[2:]
            # Remove markdown bold markers for PDF
            text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
            pdf.cell(8)  # indent
            pdf.multi_cell(0, 6, f"\u2022  {text}")
            pdf.ln(1)
        # Horizontal rule
        elif stripped in ("---", "***", "___"):
            pdf.set_draw_color(200, 200, 200)
            pdf.line(10, pdf.get_y(), 200, pdf.get_y())
            pdf.ln(5)
        # Regular paragraph
        else:
            pdf.set_font("Helvetica", "", 11)
            pdf.set_text_color(40, 40, 40)
            # Remove markdown bold markers
            text = re.sub(r"\*\*([^*]+)\*\*", r"\1", stripped)
            pdf.multi_cell(0, 6, text)
            pdf.ln(2)

    pdf.output(str(filepath))


# ── Main tool ─────────────────────────────────────────────

FORMAT_CONFIG = {
    "markdown": {"ext": ".md", "writer": _write_markdown},
    "text":     {"ext": ".txt", "writer": _write_text},
    "html":     {"ext": ".html", "writer": _write_html},
    "docx":     {"ext": ".docx", "writer": _write_docx},
    "xlsx":     {"ext": ".xlsx", "writer": _write_xlsx},
    "pdf":      {"ext": ".pdf", "writer": _write_pdf},
}

# Aliases so the LLM can say "word", "doc", "excel", "spreadsheet", etc.
FORMAT_ALIASES = {
    "doc": "docx", "word": "docx",
    "excel": "xlsx", "spreadsheet": "xlsx", "csv": "xlsx",
    "md": "markdown",
    "txt": "text",
}


@tool
async def generate_document(
    content: str,
    title: str = "Document",
    format: str = "markdown",
) -> str:
    """Generate a downloadable document file from the given content.
    Creates the file and returns a download link.

    Args:
        content: The full text content for the document.
        title: The title/filename for the document (without extension).
        format: Output format — 'markdown', 'text', 'html', 'docx', 'xlsx', or 'pdf'. Also accepts aliases like 'word', 'doc', 'excel'. Defaults to 'markdown'.
    """
    # Resolve aliases
    fmt = format.lower().strip()
    fmt = FORMAT_ALIASES.get(fmt, fmt)

    if fmt not in FORMAT_CONFIG:
        return f"Error: Unsupported format '{format}'. Supported: {', '.join(FORMAT_CONFIG.keys())}"

    logger.info("tool.generate_document", title=title, format=fmt, content_length=len(content))

    config = FORMAT_CONFIG[fmt]

    # Clean filename
    safe_title = "".join(c if c.isalnum() or c in " -_" else "_" for c in title).strip()
    if not safe_title:
        safe_title = "document"

    # Generate unique file ID
    file_id = uuid.uuid4().hex[:16]
    filename = f"{safe_title}{config['ext']}"
    filepath = DOCS_DIR / f"{file_id}_{filename}"

    # Write document using the format-specific writer
    try:
        config["writer"](content, filepath, title)
    except Exception as e:
        logger.error("tool.generate_document.write_error", format=fmt, error=str(e))
        return f"Error generating {fmt.upper()} file: {str(e)}"

    # Track for cleanup and serving
    file_size = filepath.stat().st_size
    _generated_docs[file_id] = {
        "path": str(filepath),
        "filename": filename,
        "created_at": time.time(),
        "format": fmt,
    }

    # Cleanup old files (older than 1 hour)
    _cleanup_old_docs()

    api_url = os.environ.get("RENDER_EXTERNAL_URL", "")
    if not api_url:
        api_url = os.environ.get("API_URL", "http://localhost:8000")

    download_url = f"{api_url}/api/v1/documents/download/{file_id}"

    size_str = f"{file_size:,} bytes" if file_size < 1024 else f"{file_size / 1024:.1f} KB"

    return (
        f"Document generated successfully!\n\n"
        f"**Title:** {title}\n"
        f"**Format:** {fmt.upper()}\n"
        f"**Size:** {size_str}\n\n"
        f"Download link: {download_url}\n\n"
        f"[DOWNLOAD_LINK:{file_id}:{filename}]"
    )


def get_generated_doc(file_id: str) -> dict | None:
    """Retrieve a generated document by its file ID."""
    return _generated_docs.get(file_id)


def _cleanup_old_docs():
    """Remove documents older than 1 hour."""
    now = time.time()
    expired = [fid for fid, info in _generated_docs.items()
               if now - info["created_at"] > 3600]
    for fid in expired:
        info = _generated_docs.pop(fid)
        try:
            Path(info["path"]).unlink(missing_ok=True)
        except Exception:
            pass
    if expired:
        logger.info("docs.cleanup", removed=len(expired))


# Tool collections for document agents
DOC_READER_TOOLS = [analyze_document, extract_key_points]
DOC_GENERATOR_TOOLS = [generate_document]
